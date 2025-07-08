import logging
import io
import os
import re
from django.conf import settings

logger = logging.getLogger(__name__)

# Import necessary packages for different file types
try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False

try:
    import pandas as pd
    PANDAS_SUPPORT = True
except ImportError:
    PANDAS_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

def extract_text_from_pdf(document):
    """Legacy function for backward compatibility"""
    try:
        if not PDF_SUPPORT:
            return False
            
        document.file.seek(0)
        file_content = document.file.read()
        document.file.seek(0)
        
        pdf = PdfReader(io.BytesIO(file_content))
        document.page_count = len(pdf.pages)
        
        all_text = ""
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                all_text += page_text + "\n\n"
        
        document.extracted_text = all_text.strip()
        document.is_processed = True
        document.save()
        
        return True
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF {document.title}: {str(e)}")
        document.is_processed = False
        document.extracted_text = f"[Error extracting text: {str(e)}]"
        document.save()
        return False

def extract_text_from_file(file_obj, page_range_str: str = None, filename: str = None):
    """
    Extract text from a file (PDF, TXT) with optional page range for PDFs.

    Args:
        file_obj: File-like object (e.g., from request.FILES)
        page_range_str: Optional string specifying page ranges (e.g., "1-5,7,10-15")
        filename: Optional filename to help determine file type if file_obj lacks one.

    Returns:
        str: Extracted text or an error message.
    """
    try:
        # Determine file type from filename if available
        if filename:
            file_extension = os.path.splitext(filename)[1].lower()
        else:
            file_extension = "" # Default if no filename

        # Reset file position
        file_obj.seek(0)
        
        # Handle PDF
        if file_extension == '.pdf' or (hasattr(file_obj, 'name') and file_obj.name.lower().endswith('.pdf')):
            pdf_reader = PdfReader(file_obj)
            text = ""
            
            # Parse page ranges if provided
            if page_range_str:
                selected_pages = set()
                ranges = page_range_str.split(',')
                for r in ranges:
                    if '-' in r:
                        try:
                            start, end = map(int, r.split('-'))
                            # Adjust for 0-based indexing and include the end page
                            selected_pages.update(range(start - 1, end))
                        except ValueError:
                            pass # Ignore invalid ranges
                    else:
                        try:
                            # Adjust for 0-based indexing
                            selected_pages.add(int(r) - 1)
                        except ValueError:
                            pass # Ignore invalid page numbers
                
                # Extract text from selected pages
                for page_num in sorted(list(selected_pages)):
                    if 0 <= page_num < len(pdf_reader.pages):
                        text += pdf_reader.pages[page_num].extract_text() or ""
            else:
                # Extract from all pages if no range is specified
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
            
            return text
            
        # Handle TXT files
        elif file_extension == '.txt' or (hasattr(file_obj, 'name') and file_obj.name.lower().endswith('.txt')):
            return file_obj.read().decode('utf-8')
            
        else:
            return f"[Error: Unsupported file type with extension '{file_extension}']"

    except Exception as e:
        return f"[Error extracting text: {str(e)}]"

def _extract_text_from_pdf_content(file_content, page_ranges=None):
    """
    Extract text from PDF content, optionally from specific pages.
    
    Args:
        file_content: Binary content of the PDF file
        page_ranges: Optional list of page ranges to extract. Format:
                     - [1, 5, 10] - Extract pages 1, 5, and 10 (1-indexed)
                     - [(1, 5), 10, (20, 30)] - Extract pages 1-5, 10, and 20-30
    
    Returns:
        str: Extracted text
    """
    if not PDF_SUPPORT:
        logger.error("PDF support not available. Please install PyPDF2")
        return "[PDF support not available. Install PyPDF2]"
        
    try:
        logger.info("Creating PDF reader from file content")
        pdf = PdfReader(io.BytesIO(file_content))
        total_pages = len(pdf.pages)
        logger.info(f"PDF has {total_pages} pages")
        text = ""
        
        # Parse page ranges if it's a string
        if isinstance(page_ranges, str):
            logger.info(f"Parsing page ranges string: {page_ranges}")
            page_ranges = _parse_page_ranges_str(page_ranges)
            logger.info(f"Parsed page ranges: {page_ranges}")
            
        # Process page ranges if provided
        if page_ranges:
            pages_to_extract = []
            
            for page_range in page_ranges:
                if isinstance(page_range, tuple) or (isinstance(page_range, list) and len(page_range) == 2):
                    # Handle page range (start, end)
                    start, end = page_range
                    # Convert from 1-indexed (user-friendly) to 0-indexed (internal)
                    start = max(0, start - 1)  # Ensure not negative
                    end = min(total_pages, end)  # Ensure not beyond total pages
                    pages_to_extract.extend(range(start, end))
                else:
                    # Handle single page
                    page_num = page_range - 1  # Convert from 1-indexed to 0-indexed
                    if 0 <= page_num < total_pages:
                        pages_to_extract.append(page_num)
            
            # Remove duplicates and sort
            pages_to_extract = sorted(set(pages_to_extract))
            
            # Log which pages are being extracted
            logger.info(f"Extracting text from specific pages: {[p+1 for p in pages_to_extract]} (1-indexed)")
            logger.info(f"Page ranges requested: {page_ranges}")
            
            if not pages_to_extract:
                logger.warning("No valid pages to extract based on provided page ranges")
                return "[No valid pages to extract based on provided page ranges]"
            
            # Extract text from selected pages
            for page_num in pages_to_extract:
                try:
                    logger.info(f"Extracting text from page {page_num + 1}")
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        # Add clear page boundary markers
                        text += f"==================== PAGE {page_num + 1} ====================\n"
                        text += page_text + "\n"
                        text += f"==================== END OF PAGE {page_num + 1} ====================\n\n"
                        logger.info(f"Successfully extracted {len(page_text)} characters from page {page_num + 1}")
                    else:
                        logger.warning(f"No text extracted from page {page_num + 1}")
                except Exception as e:
                    logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    text += f"[Error extracting text from page {page_num + 1}: {str(e)}]\n"
        else:
            # Extract text from all pages (original behavior)
            logger.info(f"Extracting text from all {total_pages} pages")
            for page_num in range(total_pages):
                try:
                    logger.info(f"Extracting text from page {page_num + 1}")
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        # Add page markers even when extracting all pages
                        text += f"==================== PAGE {page_num + 1} ====================\n"
                        text += page_text + "\n"
                        text += f"==================== END OF PAGE {page_num + 1} ====================\n\n"
                        logger.info(f"Successfully extracted {len(page_text)} characters from page {page_num + 1}")
                    else:
                        logger.warning(f"No text extracted from page {page_num + 1}")
                except Exception as e:
                    logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    text += f"[Error extracting text from page {page_num + 1}: {str(e)}]\n"
        
        if not text.strip():
            logger.error("No text was extracted from any page")
            return "[No text could be extracted from the PDF]"
            
        logger.info(f"Successfully extracted a total of {len(text)} characters from the PDF")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return f"[Error processing PDF: {str(e)}]"

def _extract_text_from_text_file(file_content):
    """Extract text from text-based files"""
    try:
        # Try UTF-8 first
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Fall back to latin-1
            return file_content.decode('latin-1')
        except Exception as e:
            logger.error(f"Error decoding text file: {e}")
            return f"[Error decoding text file: {str(e)}]"

def _extract_text_from_image(file_content):
    """Extract text from image files using OCR"""
    if not OCR_SUPPORT:
        return "[OCR support not available. Install pytesseract and pillow]"
        
    try:
        # Open image from binary content
        image = Image.open(io.BytesIO(file_content))
        
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(image)
        
        return text.strip() or "[No text detected in image]"
    except Exception as e:
        logger.error(f"Error extracting text from image: {e}")
        return f"[Error extracting text from image: {str(e)}]"

def _extract_text_from_docx(file_content):
    """Extract text from Word documents"""
    if not DOCX_SUPPORT:
        return "[DOCX support not available. Install python-docx]"
        
    try:
        # Save temporary file
        temp_file = io.BytesIO(file_content)
        
        # Open with python-docx
        doc = docx.Document(temp_file)
        
        # Extract text from paragraphs
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join([cell.text for cell in row.cells if cell.text]) + "\n"
        
        return text.strip() or "[No text content in document]"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return f"[Error extracting text from DOCX: {str(e)}]"

def _extract_text_from_excel(file_content):
    """Extract text from Excel files"""
    if not PANDAS_SUPPORT:
        return "[Excel support not available. Install pandas and openpyxl]"
        
    try:
        # Use pandas to read Excel
        excel_file = io.BytesIO(file_content)
        
        # Read all sheets
        excel_data = pd.read_excel(excel_file, sheet_name=None)
        
        text = ""
        
        # Process each sheet
        for sheet_name, df in excel_data.items():
            text += f"===== SHEET: {sheet_name} =====\n\n"
            
            # Convert dataframe to string
            text += df.to_string(index=False) + "\n\n"
        
        return text.strip() or "[No text content in Excel file]"
    except Exception as e:
        logger.error(f"Error extracting text from Excel: {e}")
        return f"[Error extracting text from Excel: {str(e)}]"

def _parse_page_ranges_str(page_ranges_str):
    """
    Parse a page range string into a list of page numbers and ranges.
    
    Args:
        page_ranges_str: String like "1-5,7,10-15" or single page like "3"
        
    Returns:
        List of integers and tuples, e.g. [(1,5), 7, (10,15)] or [3] for single page
    """
    if not page_ranges_str:
        return None
        
    try:
        page_ranges = []
        # Split by comma and process each range
        for range_str in page_ranges_str.split(','):
            range_str = range_str.strip()
            if '-' in range_str:
                start, end = map(int, range_str.split('-'))
                # Store as tuple for range
                page_ranges.append((start, end))
            else:
                # Store as single integer
                page_ranges.append(int(range_str))
        
        logger.info(f"Parsed page ranges '{page_ranges_str}' into: {page_ranges}")
        return page_ranges
    except ValueError as e:
        logger.error(f"Invalid page range format: {page_ranges_str}")
        return None

def extract_single_page_content(file_obj, page_number):
    """
    Extract content from a specific single page of a PDF.
    
    Args:
        file_obj: File-like object (e.g., from request.FILES)
        page_number: Single page number (1-indexed)
        
    Returns:
        str: Extracted text from the specified page or error message
    """
    try:
        if not PDF_SUPPORT:
            logger.error("PDF support not available")
            return "[PDF support not available. Install PyPDF2]"
            
        # Reset file position
        file_obj.seek(0)
        
        pdf_reader = PdfReader(file_obj)
        total_pages = len(pdf_reader.pages)
        
        # Validate page number
        if page_number < 1 or page_number > total_pages:
            logger.error(f"Invalid page number {page_number}. PDF has {total_pages} pages.")
            return f"[Invalid page number {page_number}. PDF has {total_pages} pages.]"
        
        # Convert to 0-indexed
        page_index = page_number - 1
        
        # Extract text from the specific page
        page = pdf_reader.pages[page_index]
        page_text = page.extract_text()
        
        if not page_text:
            logger.warning(f"No text found on page {page_number}")
            return f"[No text found on page {page_number}]"
        
        # Format with page markers for consistency
        formatted_text = f"==================== PAGE {page_number} ====================\n"
        formatted_text += page_text + "\n"
        formatted_text += f"==================== END OF PAGE {page_number} ====================\n"
        
        logger.info(f"Successfully extracted {len(page_text)} characters from page {page_number}")
        return formatted_text
        
    except Exception as e:
        logger.error(f"Error extracting text from page {page_number}: {str(e)}")
        return f"[Error extracting text from page {page_number}: {str(e)}]"

def validate_page_range(page_ranges_str, total_pages):
    """
    Validate that page ranges are within the document's page count.
    
    Args:
        page_ranges_str: String like "1-5,7,10-15" or single page like "3"
        total_pages: Total number of pages in the document
        
    Returns:
        dict: {'valid': bool, 'message': str, 'adjusted_ranges': str}
    """
    if not page_ranges_str:
        return {'valid': True, 'message': 'No page range specified', 'adjusted_ranges': None}
    
    try:
        page_ranges = _parse_page_ranges_str(page_ranges_str)
        if not page_ranges:
            return {'valid': False, 'message': 'Invalid page range format', 'adjusted_ranges': None}
        
        invalid_pages = []
        valid_ranges = []
        
        for page_range in page_ranges:
            if isinstance(page_range, tuple):
                start, end = page_range
                if start < 1 or end > total_pages:
                    invalid_pages.append(f"{start}-{end}")
                else:
                    valid_ranges.append(f"{start}-{end}")
            else:
                # Single page
                if page_range < 1 or page_range > total_pages:
                    invalid_pages.append(str(page_range))
                else:
                    valid_ranges.append(str(page_range))
        
        if invalid_pages:
            message = f"Invalid pages for document with {total_pages} pages: {', '.join(invalid_pages)}"
            if valid_ranges:
                message += f". Valid ranges: {', '.join(valid_ranges)}"
            return {
                'valid': False, 
                'message': message,
                'adjusted_ranges': ','.join(valid_ranges) if valid_ranges else None
            }
        
        return {
            'valid': True,
            'message': f"All pages valid for document with {total_pages} pages",
            'adjusted_ranges': page_ranges_str
        }
        
    except Exception as e:
        return {
            'valid': False,
            'message': f"Error validating page ranges: {str(e)}",
            'adjusted_ranges': None
        }
